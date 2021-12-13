from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Mm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
import sys

#func to insert text from table column to other table column
def insert_column_in_table(to_table_column, from_table_column, n, ignore):
    for i in range(n):
        ch = 0
        #print(repr(from_table_column.cells[i].text))
        if ignore:
            to_table_column.cells[i].text = from_table_column.cells[i].text
            continue
        for line in from_table_column.cells[i].text.split('\n'): 
            if line != "":
                if ch == 1:
                    to_table_column.cells[i].add_paragraph(line, style='List Bullet')
                else:
                    to_table_column.cells[i].text = line
                    ch+=1

#insert time period from console
time_period_start, time_period_stop, month = 0, 0, 0
try:
    time_period_start, time_period_stop, month = int(sys.argv[1]), int(sys.argv[2]), int(sys.argv[3])
    print(time_period_start, time_period_stop)
except ValueError:
    print("Check linux attributes")
    sys.exit()
except IndexError:
    print("put start, stop period")
    sys.exit()


#get template of table, dont forget name file
doc_template = Document("template_table.docx").tables[0]

#get size of table
rows = len(doc_template.rows)
columns = len(doc_template.columns)

#get header of table, fucking structure
header_column_number = doc_template.columns[0]
header_column_text = doc_template.columns[1]

#create new document
doc_example = Document()
section = doc_example.sections[-1]
#table.style = doc_example.styles["Table Grid"]
#set orientation
section.orientation = WD_ORIENT.LANDSCAPE

#set width height
section.page_width = Inches(11.69)
section.page_height = Inches(8.27)

#create table in new document
table_from_example = doc_example.add_table(rows=rows, cols=columns+7)

table_from_example.style = doc_example.styles["Table Grid"]
table_from_example.autofit = True
styles = table_from_example.style
font = styles.font
font.size = Pt(10)
font.name = 'Times New Roman'

#insert fucking header
insert_column_in_table(table_from_example.columns[0],header_column_number, rows, 1)
insert_column_in_table(table_from_example.columns[1],header_column_text, rows, 0)

#set size default header
table_from_example.columns[0].width = Mm(15)
table_from_example.columns[1].width = Mm(50)

#set size other blocks
for i in range(2,len(table_from_example.columns)):
    table_from_example.columns[i].width = Mm(10)

# i dont know, but its not working.
for column in table_from_example.columns:
    for cell in column.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

#set date in column
for i in range(2,len(doc_example.tables[0].columns)):
    doc_example.tables[0].columns[i].cells[0].text = "{:0>2}".format(time_period_start)
    time_period_start += 1

#set 1 column
for cell in doc_example.tables[0].column_cells(2):
    cell.text = "1"

#table_from_example.style = doc_template.styles['Table Grid']

#font = styles.font
#font.size = Pt(10)
#font.name = 'Arial'
#save document
doc_example.save("test.docx")

